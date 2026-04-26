"""Unit tests for the document_reader module in ecss_mcp_server."""

# Import Python libraries
import logging

# Import third-party libraries
import docx.document
import pytest

# Import local modules
import ecss_mcp_server.document_reader

logger = logging.getLogger(__name__)

def test_get_doc_ids():
    """Test the get_doc_ids function to ensure it returns a list of strings."""
    doc_ids = ecss_mcp_server.document_reader.get_doc_ids()
    # Assert that the returned value is a list
    assert isinstance(doc_ids, list)
    # Assert that all items in the list are strings
    assert all(isinstance(doc_id, str) for doc_id in doc_ids)

def test_load_document():
    """Test the load_document function to ensure it loads all documents in the document library correctly."""
    # Get the list of document IDs to test loading documents
    doc_ids = ecss_mcp_server.document_reader.get_doc_ids()
    # If there are no documents available, skip the test
    if not doc_ids:
        pytest.skip("No documents available to test load_document")
    for doc_id in doc_ids:
        doc = ecss_mcp_server.document_reader.load_document(doc_id)
        # Assert that the returned object is an instance of docx.document.Document
        assert isinstance(doc, docx.document.Document), (
            f"load_document('{doc_id}') did not return a docx.document.Document instance"
        )
        assert isinstance(doc.paragraphs, list), (
            f"Document '{doc_id}' has no accessible paragraphs attribute"
        )
        assert any(p.text.strip() for p in doc.paragraphs) or len(doc.tables) > 0, (
            f"Document '{doc_id}' appears to be empty, no text or tables found"
        )

def test_parse_toc_line():
    """Test parse_toc_line correctly parses valid TOC lines and raises errors for invalid lines."""
    # Test valid TOC line with section number
    toc_line1 = "1\tIntroduction\t1"
    toc_entry1 = ecss_mcp_server.document_reader.parse_toc_line(toc_line1)
    assert toc_entry1.section_number == "1"
    assert toc_entry1.heading_text == "Introduction"
    assert toc_entry1.page_number == 1

    # Test valid TOC line with annex section
    toc_line2 = "Annex A\tAdditional Information\t10"
    toc_entry2 = ecss_mcp_server.document_reader.parse_toc_line(toc_line2)
    assert toc_entry2.section_number == "Annex A"
    assert toc_entry2.heading_text == "Additional Information"
    assert toc_entry2.page_number == 10  # noqa: PLR2004

    # Test valid TOC line with no section number
    toc_line3 = "Bibliography\t20"
    toc_entry3 = ecss_mcp_server.document_reader.parse_toc_line(toc_line3)
    assert toc_entry3.section_number is None
    assert toc_entry3.heading_text == "Bibliography"
    assert toc_entry3.page_number == 20  # noqa: PLR2004

    # Test invalid TOC line that does not match expected formats
    invalid_toc_line = "This is not a valid TOC entry"
    with pytest.raises(ValueError, match="does not match expected formats"):
        ecss_mcp_server.document_reader.parse_toc_line(invalid_toc_line)

    # Test numeric section number with spaces but no tab/page number
    toc_line4 = "1.1.1  Standards"
    toc_entry4 = ecss_mcp_server.document_reader.parse_toc_line(toc_line4)
    assert toc_entry4.section_number == "1.1.1"
    assert toc_entry4.heading_text == "Standards"
    assert toc_entry4.page_number is None

    # Test that bare label "TOC" (no section, no tab) still raises ValueError
    with pytest.raises(ValueError, match="does not match expected formats"):
        ecss_mcp_server.document_reader.parse_toc_line("TOC")

    # Test that an empty string still raises ValueError
    with pytest.raises(ValueError, match="does not match expected formats"):
        ecss_mcp_server.document_reader.parse_toc_line("")

def test_extract_toc(subtests):
    """Test the extract_toc function for all documents in the document library."""
    for j, doc_id in enumerate(ecss_mcp_server.document_reader.get_doc_ids()):
        logger.info("Testing extract_toc on document '%s'", doc_id)
        with subtests.test(doc_id=doc_id, index=j):
            doc = ecss_mcp_server.document_reader.load_document(doc_id)
            # Calling extract_toc on the document should return a list of TOCEntry objects without errors
            toc = ecss_mcp_server.document_reader.extract_toc(doc)
            # Assert all TOCEntry values have a paragraph_index (i.e. each heading was located in the document body)
            for toc_entry in toc:
                assert toc_entry.paragraph_index is not None, (
                    f"TOC entry '{toc_entry.heading_text}' in document '{doc_id}' has paragraph_index None"
                )
            # Assert paragraph_index values are strictly increasing (headings are matched in order)
            paragraph_indices = [toc_entry.paragraph_index for toc_entry in toc]
            for i, paragraph_index in enumerate(paragraph_indices):
                if i == 0:
                    continue
                assert paragraph_index > paragraph_indices[i - 1], (
                    f"paragraph_index not strictly increasing at '{toc[i].heading_text}' in document '{doc_id}': "
                    f"{paragraph_indices[i - 1]} -> {paragraph_index}"
                )

