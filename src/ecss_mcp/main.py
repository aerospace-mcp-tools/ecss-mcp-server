"""FastMCP server exposing ECSS standards documents as MCP tools."""

# Import Python libraries
from pathlib import Path

# Import third-party libraries
from docx import Document
from fastmcp import FastMCP

app = FastMCP("ecss-mcp-server")

# Load document from document library using document ID
def load_document(doc_id: str) -> Document:
    """
    Load a document from the document library using its document ID.

    Args:
        doc_id (str): The document ID of the document to load.

    Returns:
        Document: The loaded document object.

    """
    doc_path = Path(f"/app/documents/{doc_id}.docx")
    if not doc_path.exists():
        msg = f"Document with ID {doc_id} not found in document folder."
        raise FileNotFoundError(msg)
    return Document(str(doc_path))

# Extract the Table of Contents from the document
def extract_toc(document: Document) -> list:  # noqa: C901
    """
    Extract the Table of Contents from a document.

    Args:
        document (Document): The document object to extract the ToC from.

    Returns:
        list: A list of table of contents entries, each containing the section number,
            heading text, page number and paragraph index.

    """

    class TocEntry:
        def __init__(self) -> None:
            self.section = None # Section number
            self.heading = None # Heading text
            self.page = None # Page number
            self.paragraph_i = None # Paragraph index

    # Table of Contents
    toc_styles = {'toc 1', 'toc 2', 'toc 3'}
    toc_paragraphs = [p for p in document.paragraphs if p.style.name.lower() in toc_styles]

    tocs = []

    for p in toc_paragraphs:
        toc_list = list(p.text.strip())
        section = []
        page = []
        text = []
        # Extract leading numbers (section numbers)
        for i, char in enumerate(toc_list):
            if char.isdigit() or char == '.':
                section.append(char)
            else:
                text = toc_list[i:]
                break
        # Extract page number from the end
        for i, char in enumerate(reversed(text)):
            if char.isdigit():
                page.append(char)
            else:
                text = text[:-i]  # Remove page number from text
                page.reverse()
                break
        toc_entry = TocEntry()
        toc_entry.section = ''.join(section).strip()
        # Extract section heading text
        toc_entry.heading = ''.join(text).strip()
        toc_entry.page = ''.join(page).strip()
        tocs.append(toc_entry)

    # Find heading paragraphs that match the table of contents by iterating through all headings
    heading_styles = {'heading 0', 'heading 1', 'heading 2', 'heading 3', 'heading 4', 'heading 5'}
    for i, p in enumerate(document.paragraphs):
        if p.style.name.lower() in heading_styles:
            heading_text = p.text.strip()
            # Match the heading text with the next unmatched ToC entry
            for toc in tocs:
                if toc.heading == heading_text and toc.paragraph_i is None:
                    toc.paragraph_i = i
                    break
    return tocs

# Extract the List of Figures or Tables from the document
def extract_fots(document: Document) -> list:
    """
    Extract the List of Figures or Tables from a document.

    Args:
        document (Document): The document object to extract the List of Figures or Tables from.

    Returns:
        list: A list of figures or tables, each containing the type (figure or table), number, text, and page number.

    """
    # List of figures and tables
    list_styles = {'table of figures'}
    list_paragraphs = [p for p in document.paragraphs if p.style.name.lower() in list_styles]

    # Defne classes for figures and tables
    class Fot:
        def __init__(self, element: str, number: str, text: str, page: str) -> None:
            self.element = element
            self.number = number
            self.text = text
            self.page = page

    fots = []

    for p in list_paragraphs:
        list_line = list(p.text.strip())
        element = ''
        number = []
        page = []
        # Determine if it's a figure or table
        if p.text.strip().startswith('Figure'):
            element = 'Figure'
            text = list_line[6:]  # Remove 'Figure' from text
        elif p.text.strip().startswith('Table'):
            element = 'Table'
            text = list_line[5:]  # Remove 'Table' from text
        else:
            continue  # Skip if it doesn't start with Figure or Table

        # Extract figure number from the start
        # Strip leading spaces
        text = ''.join(text).strip()
        text = list(text)
        for i, char in enumerate(text):
            if char != ' ':
                number.append(char)
            else:
                text = text[i:]
                break

        # Extract page number from the end
        for i, char in enumerate(reversed(text)):
            if char.isdigit():
                page.append(char)
            else:
                if i > 0:
                    text = text[:-i]  # Remove page number from text
                page.reverse()
                break
        page = ''.join(page).strip()
        text = ''.join(text).strip()
        number = ''.join(number).strip()
        fot_entry = Fot(element, number, text, page)
        fots.append(fot_entry)
    return fots

def extract_section(document: Document, section: str, heading: str) -> str:
    """
    Extract the text of a specific section from a document.

    Args:
        document (Document): The document object to extract the section from.
        section (str): The section to extract as taken from table of contents (i.e. 1.2)
        heading (str): The heading text of the section to extract as taken from the
            table of contents (e.g. "Thermal Analysis").

    Returns:
        str: The text of the specified section.

    """
    toc = extract_toc(document)
    # Match the ToC entry to the section number and heading
    matched_toc_entry = next(
        (toc_entry for toc_entry in toc if toc_entry.section == section and toc_entry.heading == heading),
        None,
    )
    if matched_toc_entry is None:
        msg = "Section not found in table of contents"
        raise ValueError(msg)
    if matched_toc_entry.paragraph_i is None:
        msg = "Section could not be located in the document body"
        raise ValueError(msg)
    # Find next ToC entry at the same level as this, if not default to end of document
    next_toc_entry = toc[-1]
    for toc_entry in toc:
        if (
            toc_entry.paragraph_i is not None
            and len(toc_entry.section) == len(matched_toc_entry.section)
            and toc_entry.paragraph_i > matched_toc_entry.paragraph_i
        ):
            next_toc_entry = toc_entry
            break
    # Iterate between paragraphs from the matched ToC entry to the next ToC entry at the same level and extract text
    section_text = ""
    for p in document.paragraphs[matched_toc_entry.paragraph_i : next_toc_entry.paragraph_i]:
        section_text += p.text + "\n"
    return section_text


@app.tool()
def get_doc_ids() -> list:
    """
    Get a list of document IDs from the document library.

    Returns:
        list: A list of all documents in the document library.

    """
    return [doc_file.stem for doc_file in Path("/app/documents").glob("*.docx")]

@app.tool()
def get_doc_summary(doc_id: str) -> str:
    """
    Get a summary of a document given its ID.

    Integration pattern: Only use this tool after using the get_doc_ids tool to get a list of document IDs in
    the document library, and only use it with those document IDs.

    Args:
        doc_id (str): The ECSS ID of the document to summarize.

    Returns:
        str: A summary of the document including document scope and table of contents.

    """
    doc = load_document(doc_id)
    tocs = extract_toc(doc)
    fots = extract_fots(doc)
    scope = extract_section(doc, "1", "Scope")
    summary = f"Document ID: {doc_id}\n\nSummary:\n"
    summary += f"Scope:\n{scope}\n\n"
    summary += "Table of Contents:\n"
    summary += "'section': 'heading'\n"
    for toc in tocs:
        summary += f"'{toc.section}': '{toc.heading}'\n"
    summary += "\nList of Figures and Tables:\n"
    summary += "'element': 'number' 'text'\n"
    for fot in fots:
        summary += f"'{fot.element}': '{fot.number}' '{fot.text}'\n"
    return summary

@app.tool()
def get_section_text(doc_id: str, section: str, heading: str) -> str:
    """
    Get the text of a specific section from a document given its section number and heading.

    Integration pattern: Only use this tool after using the get_doc_summary tool to get the table of contents
    for a document, and only use it with section numbers and headings that are present in the table of
    contents for that document.

    Args:
        doc_id (str): The ECSS ID of the document to extract the section from.
        section (str): The section number to extract as taken from table of contents (e.g. "5.5.3")
        heading (str): The heading text of the section to extract as taken from the table of contents
            (e.g. "Thermal Analysis").

    Returns:
        str: The text of the specified section.

    """
    doc = load_document(doc_id)
    return extract_section(doc, section, heading)

def main() -> None:
    """Entry point for the ecss-mcp-server script."""
    app.run()

if __name__ == "__main__":
    main()
